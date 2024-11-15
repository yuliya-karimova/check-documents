from docx import Document
from docx.shared import Pt, Inches
import markdown
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import tempfile
import base64
from io import BytesIO
from PIL import Image as PILImage
from lxml import etree
from bs4 import BeautifulSoup
from reportlab.lib.units import inch
import PyPDF2
from fpdf import FPDF
import os
import openpyxl
from openpyxl import Workbook, load_workbook

def add_html_to_docx(doc, html):
    soup = BeautifulSoup(html, 'html.parser')
    for element in soup.recursiveChildGenerator():
        if element.name:
            if element.name == 'h1':
                doc.add_heading(element.text, level=1)
            elif element.name == 'h2':
                doc.add_heading(element.text, level=2)
            elif element.name == 'h3':
                doc.add_heading(element.text, level=3)
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    if li.text:
                        doc.add_paragraph(li.text, style='ListBullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    if li.text:
                        doc.add_paragraph(li.text, style='ListNumber')
            elif element.name == 'em':
                run = doc.add_paragraph().add_run(element.text)
                run.italic = True
            elif element.name == 'img':
                img_data = base64.b64decode(element['src'].split(',')[1])
                with BytesIO(img_data) as img_stream:
                    doc.add_picture(img_stream, width=Inches(6))
                    


def add_html_to_elements(html, elements):
    soup = BeautifulSoup(html, 'html.parser')

    # Регистрация шрифта DejaVuSans
    pdfmetrics.registerFont(TTFont('DejaVuSans', 'data/fonts/DejaVuSans.ttf'))
    styles = getSampleStyleSheet()
    styles['Normal'].fontName = 'DejaVuSans'
    styles['BodyText'].fontName = 'DejaVuSans'
    styles['Heading1'].fontName = 'DejaVuSans'
    styles['Heading2'].fontName = 'DejaVuSans'
    styles['Heading3'].fontName = 'DejaVuSans'
    
    for element in soup.recursiveChildGenerator():
        if element.name:
            if element.name == 'h1':
                elements.append(Paragraph(element.text, styles['Heading1']))
            elif element.name == 'h2':
                elements.append(Paragraph(element.text, styles['Heading2']))
            elif element.name == 'h3':
                elements.append(Paragraph(element.text, styles['Heading3']))
            elif element.name == 'p':
                elements.append(Paragraph(element.text, styles['BodyText']))
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    elements.append(Paragraph(f'• {li.text}', styles['BodyText']))
            elif element.name == 'ol':
                for index, li in enumerate(element.find_all('li'), 1):
                    elements.append(Paragraph(f'{index}. {li.text}', styles['BodyText']))
            elif element.name == 'a':
                elements.append(Paragraph(f'<a href="{element["href"]}">{element.text}</a>', styles['BodyText']))
            elif element.name == 'img':
                img_data = base64.b64decode(element['src'].split(',')[1])
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
                    tmpfile.write(img_data)
                    tmpfile.flush()
                    with PILImage.open(tmpfile.name) as img:
                        img_width, img_height = img.size
                        ratio = min(5*inch / img_width, 5*inch / img_height)
                        img_width = int(img_width * ratio)
                        img_height = int(img_height * ratio)
                        elements.append(Image(tmpfile.name, width=img_width, height=img_height))
                        elements.append(Spacer(1, 12))

def create_docx_file(report_text, file_path):
    doc = Document()
    html = markdown.markdown(report_text)
    add_html_to_docx(doc, html)
    doc.save(file_path)

def create_pdf_file(report_text, file_path):
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    elements = []

    html = markdown.markdown(report_text)
    add_html_to_elements(html, elements)

    doc.build(elements)

def generate_and_return_files(folder, full_text):
    # Генерация файлов (DOCX и PDF)
    docx_file_path = os.path.join(folder, "report.docx")
    pdf_file_path = os.path.join(folder, "report.pdf")

    # Создаем файлы DOCX и PDF с помощью переданного текста
    create_docx_file(full_text, docx_file_path)
    create_pdf_file(full_text, pdf_file_path)

    # Формируем URL-адреса, используя полный путь `folder`
    docx_url = f"{folder}/report.docx"
    pdf_url = f"{folder}/report.pdf"

    return docx_url, pdf_url